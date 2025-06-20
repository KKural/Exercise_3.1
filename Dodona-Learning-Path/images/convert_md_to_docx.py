from bs4 import BeautifulSoup
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx import Document
import markdown
import os
import sys

# Check for required packages and install them if they're missing
required_packages = ['python-docx', 'markdown', 'beautifulsoup4']
missing_packages = []

for package in required_packages:
    try:
        __import__(package.replace('-', '_'))
    except ImportError:
        missing_packages.append(package)

if missing_packages:
    print(f"Missing required packages: {', '.join(missing_packages)}")
    print("Please run install_requirements.py first or install them manually with:")
    print(f"pip install {' '.join(missing_packages)}")
    sys.exit(1)

# Now that we know the packages are installed, import them


def markdown_to_word(md_file, output_file):
    try:
        # Read the markdown file
        with open(md_file, 'r', encoding='utf-8') as file:
            md_content = file.read()

        print(f"Successfully read Markdown file: {md_file}")

        # Create a new Word document
        doc = Document()

        # Set font for the entire document
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Convert markdown to HTML
        html = markdown.markdown(md_content, extensions=[
                                 'tables', 'fenced_code'])

        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')

        # Process the document
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'table', 'img']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                doc.add_heading(element.text, level)

            elif element.name == 'p':
                # Check if this paragraph contains an image
                img = element.find('img')
                if img:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_src = img.get('src', '')
                    alt_text = img.get('alt', 'Image')

                    # Add image description
                    doc.add_paragraph(
                        f"[Image: {alt_text} - {img_src}]").italic = True
                else:
                    doc.add_paragraph(element.text)

            elif element.name == 'ul':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='List Bullet')

            elif element.name == 'ol':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='List Number')

            elif element.name == 'table':
                # Get number of rows and columns
                rows = element.find_all('tr')
                if not rows:
                    continue

                # Get the maximum number of cells in any row
                max_cells = max(
                    len(row.find_all(['td', 'th'])) for row in rows)

                # Create the table
                table = doc.add_table(rows=len(rows), cols=max_cells)
                table.style = 'Table Grid'

                # Fill the table
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.text.strip()

        # Save the document
        doc.save(output_file)
        return f"Document successfully saved to {output_file}"

    except Exception as e:
        return f"Error: {str(e)}"


# Usage
if __name__ == "__main__":
    input_file = "Dodona_Learning_Path_Overview.md"
    output_file = "Dodona_Learning_Path_Overview.docx"

    # Get the current directory
    current_dir = os.path.dirname(os.path.abspath(__file__))

    # Construct full paths
    input_path = os.path.join(current_dir, input_file)
    output_path = os.path.join(current_dir, output_file)

    if not os.path.exists(input_path):
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)

    print(f"Converting {input_path} to {output_path}...")
    result = markdown_to_word(input_path, output_path)
    print(result)
