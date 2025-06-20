import os
import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def markdown_to_word(md_file, output_file):
    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as file:
        md_content = file.read()

    # Create a new Word document
    doc = Document()

    # Set font for the entire document
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Convert markdown to HTML
    html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])

    # Split the HTML by headers
    sections = re.split(r'<h1>|<h2>|<h3>', html)

    # Process the title
    title = "Dodona Learning Path: Pedagogical & Technical Overview"
    doc.add_heading(title, 0)

    # Process each section
    current_section = ""
    for section in sections:
        if not section.strip():
            continue

        # Extract header and content
        if '</h1>' in section:
            header, content = section.split('</h1>', 1)
            level = 1
        elif '</h2>' in section:
            header, content = section.split('</h2>', 1)
            level = 2
        elif '</h3>' in section:
            header, content = section.split('</h3>', 1)
            level = 3
        else:
            content = section
            header = ""
            level = 0

        # Add header if it exists
        if header:
            doc.add_heading(header, level)

        # Process content - this is simplified
        # In a real implementation, you'd parse the HTML properly

        # Convert <p> paragraphs
        paragraphs = re.findall(r'<p>(.*?)</p>', content, re.DOTALL)
        for p in paragraphs:
            # Remove HTML tags
            p = re.sub(r'<.*?>', '', p)
            doc.add_paragraph(p)

        # Convert lists
        list_items = re.findall(r'<li>(.*?)</li>', content, re.DOTALL)
        for item in list_items:
            # Remove HTML tags
            item = re.sub(r'<.*?>', '', item)
            doc.add_paragraph(item, style='List Bullet')

        # Handle images - this will be placeholder as we can't access the images
        images = re.findall(r'<img.*?src="(.*?)".*?>', content)
        for img_src in images:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()

            # Extract relative path
            img_path = img_src
            # In a real implementation, you'd resolve the path properly
            # For now, add a placeholder text
            doc.add_paragraph(f"[Image: {img_path}]")

        # Add a table if present
        if '<table>' in content:
            doc.add_paragraph("Table content would be processed here")

    # Save the document
    doc.save(output_file)
    return f"Document saved to {output_file}"


# Usage
if __name__ == "__main__":
    input_file = "Dodona_Learning_Path_Overview.md"
    output_file = "Dodona_Learning_Path_Overview.docx"

    # Get the current directory
    current_dir = os.path.dirname(os.path.abspath(__file__))

    # Construct full paths
    input_path = os.path.join(current_dir, input_file)
    output_path = os.path.join(current_dir, output_file)

    result = markdown_to_word(input_path, output_path)
    print(result)
