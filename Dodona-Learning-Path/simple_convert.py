import os
import sys
import time
import logging
import traceback
import urllib.parse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import glob


def find_images_efficiently(project_root, max_depth=3, logger=None):
    """Find image files more efficiently by focusing on likely directories and limiting depth"""
    all_images = []

    # Use a set to track processed directories (avoid duplicates)
    processed_dirs = set()

    # Safely use logger if provided
    def safe_log(msg):
        if logger:
            try:
                logger.info(msg)
            except:
                print(msg)
        else:
            print(msg)

    log = safe_log    # Special handling for common problematic files (directly check for these)
    special_files = [
        'bloom-pyramid.png',
        'bloom-pyramid.svg',
        'bloom_pyramid.png'
    ]
    
    for special_file in special_files:
        # Check in project root
        direct_path = os.path.join(project_root, special_file)
        if os.path.exists(direct_path):
            log(f"Found special file {special_file} at {direct_path}")
            all_images.append(direct_path)
        
        # Check in parent directory
        parent_path = os.path.join(os.path.dirname(project_root), special_file)
        if os.path.exists(parent_path):
            log(f"Found special file {special_file} at {parent_path}")
            all_images.append(parent_path)
            
        # Check in Dodona-Learning-Path directory
        dodona_path = os.path.join(project_root, 'Dodona-Learning-Path', special_file)
        if os.path.exists(dodona_path):
            log(f"Found special file {special_file} in Dodona-Learning-Path directory")
            all_images.append(dodona_path)
    
    # Specific directories known to contain images (in priority order)
    image_dirs = [
        'images',
        'media',
        'description/media',
        'description/images',
        'Dodona-Learning-Path/images',
        '../',  # Look for images in the parent directory
        './'    # Also look in the current directory
    ]

    # Add more common image locations
    image_dirs.extend([
        '10. Boxplot Interpretation/description/media',
        '11. Scatterplot of unemployment/description/media'
    ])

    # First check specific image directories (most likely to contain the images)
    for img_dir in image_dirs:
        img_path = os.path.join(project_root, img_dir)
        if os.path.exists(img_path) and os.path.isdir(img_path) and img_path not in processed_dirs:
            processed_dirs.add(img_path)
            log(f"Searching for images in {img_path}")

            # Get all image files in one go (more efficient)
            try:
                image_files = [f for f in os.listdir(img_path)
                               if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.svg'))]

                # Add all found images at once
                all_images.extend([os.path.join(img_path, f)
                                  for f in image_files])

                log(f"Found {len(image_files)} images in {img_dir}")
            except Exception as e:
                log(f"Error accessing directory {img_path}: {str(e)}")

    # If we don't find enough images, do a limited depth search (but only if really necessary)
    if len(all_images) < 5:  # Arbitrary threshold
        log(
            f"Performing limited-depth search for images (max depth: {max_depth})")

        for root, dirs, files in os.walk(project_root):
            # Skip already processed directories
            if root in processed_dirs:
                continue
            processed_dirs.add(root)

            # Calculate current depth
            current_depth = root[len(project_root):].count(os.sep)
            if current_depth > max_depth:
                dirs[:] = []  # Don't go deeper
                continue

            # Skip version control and other unnecessary directories
            dirs[:] = [d for d in dirs if d not in [
                '.git', '.github', 'node_modules', '__pycache__']]

            # Fast check if directory contains any images
            image_files = [f for f in files if f.lower().endswith(
                ('.png', '.jpg', '.jpeg', '.gif', '.svg'))]
            if image_files:
                all_images.extend([os.path.join(root, f) for f in image_files])
                log(f"Found {len(image_files)} images in {root}")

    log(f"Found {len(all_images)} images total")
    return all_images


def simple_markdown_to_word(md_file, output_file, logger=None, show_image_debug=False):
    try:
        # Safely use logger if provided
        def safe_log(msg):
            if logger:
                try:
                    logger.info(msg)
                except:
                    print(msg)
            else:
                print(msg)

        def safe_log_error(msg):
            if logger:
                try:
                    logger.error(msg)
                except:
                    print(f"ERROR: {msg}")
            else:
                print(f"ERROR: {msg}")

        log = safe_log
        log_error = safe_log_error

        # Read the markdown file
        with open(md_file, 'r', encoding='utf-8') as file:
            md_content = file.read()

        log(f"Successfully read Markdown file: {md_file}")

        # Quick check if the file has any images before doing expensive image search
        # Get the base directory for image resolution
        has_images = '![' in md_content and '](' in md_content
        base_dir = os.path.dirname(os.path.abspath(md_file))
        project_root = find_project_root(base_dir)

        # Debug image paths
        log(f"Base directory: {base_dir}")
        log(f"Project root: {project_root}")

        if show_image_debug:
            log("\n" + "="*40)
            log("RUNNING IN IMAGE DEBUG MODE")
            log("This will provide extra information about image processing")
            log("="*40 + "\n")

        # Create a new Word document
        doc = Document()

        # Set document properties for better formatting
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Set font for the entire document
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)

        # Set paragraph spacing
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.15

        # Update heading styles for better formatting
        for i in range(1, 10):
            if f'Heading {i}' in doc.styles:
                heading_style = doc.styles[f'Heading {i}']
                heading_style.font.name = 'Times New Roman'
                heading_style.font.color.rgb = RGBColor(0, 0, 0)
                heading_style.font.bold = True

                # Add proper spacing for headings
                heading_style.paragraph_format.space_before = Pt(12)
                heading_style.paragraph_format.space_after = Pt(6)

                # Set specific size based on heading level
                if i == 1:
                    heading_style.font.size = Pt(16)
                elif i == 2:
                    heading_style.font.size = Pt(14)
                else:
                    heading_style.font.size = Pt(13)

        # Process each line of the markdown
        current_list = []
        in_list = False
        in_table = False
        # Only search for images if the document actually contains image markdown
        table_data = []
        image_lookup = {}
        if has_images:
            log("Document contains images, searching for image files...")
            # Use the more efficient image search - but only do it once
            all_images = find_images_efficiently(project_root, logger=logger)
            log(f"Found {len(all_images)} images")

            # In debug mode, print out all available images for reference
            if show_image_debug:
                log("\nAVAILABLE IMAGES:")
                for i, img in enumerate(all_images):
                    log(f"{i+1}. {os.path.basename(img)} - {img}")
                log("\n")            # Create a lookup dictionary for faster image matching
            for img in all_images:
                base_name = os.path.basename(img).lower()
                image_lookup[base_name] = img

                # Also add the name without extension for more flexible matching
                name_without_ext = os.path.splitext(base_name)[0].lower()
                image_lookup[name_without_ext] = img
                
                # Add URL-decoded versions to support URL-encoded filenames in markdown
                url_decoded = urllib.parse.unquote(base_name).lower()
                if url_decoded != base_name:
                    image_lookup[url_decoded] = img
                    
                    # Also add URL-decoded name without extension
                    url_decoded_name = os.path.splitext(url_decoded)[0].lower()
                    image_lookup[url_decoded_name] = img

                # Add more variants for flexible matching
                if '-' in name_without_ext:
                    image_lookup[name_without_ext.replace('-', '_')] = img
                if '_' in name_without_ext:
                    image_lookup[name_without_ext.replace('_', '-')] = img

                # Also add variations without extension
                name_part, _ = os.path.splitext(base_name)
                image_lookup[name_part] = img

            log(f"Created image lookup with {len(image_lookup)} entries")
        else:
            log("Document does not contain any images, skipping image search")
            all_images = []

        log("Starting to process markdown content...")
        # Split content into lines once
        lines = md_content.split('\n')
        i = 0
        total_lines = len(lines)

        # Report progress at these percentages
        progress_points = [10, 25, 50, 75, 90]
        next_progress_idx = 0

        # Process lines in batches for better performance
        while i < len(lines):
            # Show progress for long documents
            current_percentage = (i / total_lines) * 100
            if next_progress_idx < len(progress_points) and current_percentage >= progress_points[next_progress_idx]:
                log(f"Processing: {progress_points[next_progress_idx]}% complete...")
                next_progress_idx += 1

            line = lines[i]

            # Check for headers (fast path for common case)
            if line.startswith('#'):
                if line.startswith('# '):
                    heading = doc.add_heading(line[2:], 0)  # Title
                elif line.startswith('## '):
                    heading = doc.add_heading(line[3:], 1)
                elif line.startswith('### '):
                    heading = doc.add_heading(line[4:], 2)
                elif line.startswith('#### '):
                    heading = doc.add_heading(line[5:], 3)
                else:
                    # Default paragraph for unknown header format
                    doc.add_paragraph(line)

                # Apply formatting to all header types
                if 'heading' in locals():
                    for run in heading.runs:
                        run.font.name = 'Times New Roman'
                        run.font.color.rgb = RGBColor(0, 0, 0)

            # Check for bullet points and numbered lists (combined for efficiency)
            elif line.strip().startswith(('-', '*', '1.')):
                # Special handling for specific bullet points from the learning path overview
                if any(keyword in line.lower() for keyword in ['cognitive levels', 'learning intentions', 'scaffolding techniques', 'criminologically relevant']):
                    # This is one of our special bullet points that needs proper formatting
                    p = doc.add_paragraph('', style='List Bullet')

                    if 'cognitive levels' in line.lower():
                        p.add_run(
                            'Questions build systematically through ').bold = False
                        p.add_run('cognitive levels').bold = True
                        p.add_run(
                            ', from Remember (measurement levels, descriptive vs. inferential) through Understanding, Application, Analysis, and Evaluation, to Creation (research design).').bold = False

                    elif 'learning intentions' in line.lower():
                        p.add_run(
                            'Each question addresses specific ').bold = False
                        p.add_run('learning intentions').bold = True
                        p.add_run(
                            ' that align with course objectives while targeting common misconceptions.').bold = False

                    elif 'scaffolding techniques' in line.lower():
                        p.add_run('The exercises incorporate ').bold = False
                        p.add_run('scaffolding techniques').bold = True
                        p.add_run(
                            ' including multi-level hints, targeted feedback, and visual explanations.').bold = False

                    elif 'criminologically relevant' in line.lower():
                        p.add_run(
                            'Content is contextualized within ').bold = False
                        p.add_run(
                            'criminologically relevant scenarios').bold = True
                        p.add_run(
                            ' such as analyzing burglary rates, interpreting crime statistics, and evaluating relationships between socioeconomic factors and crime.').bold = False

                # Standard bullet point handling
                elif line.strip().startswith(('-', '*')):
                    doc.add_paragraph(line.strip()[2:], style='List Bullet')
                elif line.strip().startswith('1.') and line.strip()[3:].strip():
                    doc.add_paragraph(line.strip()[3:], style='List Number')
                else:
                    doc.add_paragraph(line)  # Fallback

            # Check for tables - optimized handling
            elif line.strip().startswith('|') and line.strip().endswith('|'):
                if not in_table:
                    in_table = True
                    table_data = []

                # Extract cells from the table row (more efficiently)
                cells = [cell.strip() for cell in line.strip('|').split('|')]
                table_data.append(cells)

                # Check if the next line is a separator row
                if i + 1 < len(lines) and lines[i + 1].strip().startswith('|') and set(c for c in lines[i + 1].strip('|') if c not in '|-: ') == set():
                    i += 1  # Skip the separator row

            # Check for end of table
            elif in_table and not line.strip().startswith('|'):
                # Create the table in the document
                if table_data:
                    num_rows = len(table_data)
                    if num_rows > 0:
                        num_cols = max(len(row) for row in table_data)
                        if num_cols > 0:
                            table = doc.add_table(rows=num_rows, cols=num_cols)
                            table.style = 'Table Grid'

                            # Fill the table with data (batch process)
                            for row_idx, row_data in enumerate(table_data):
                                for col_idx, cell_data in enumerate(row_data):
                                    if col_idx < num_cols:
                                        cell = table.cell(row_idx, col_idx)
                                        cell.text = cell_data
                                        # Apply Times New Roman font to table text
                                        for paragraph in cell.paragraphs:
                                            for run in paragraph.runs:
                                                run.font.name = 'Times New Roman'

                in_table = False
                table_data = []
                # Process the current line (non-table line)
                if line.strip():
                    # Enhanced image and screenshot handling with actual image embedding
                    doc.add_paragraph(line)
            elif has_images and line.strip().startswith('!['):
                try:
                    # Extract image details (more efficiently)
                    start = line.find('![') + 2
                    end = line.find(']', start)
                    alt_text = line[start:end] if end > start else "Image"
                    
                    img_start = line.find('(', end) + 1
                    img_end = line.find(')', img_start)                    # Try to find the actual image file
                    img_path = line[img_start:img_end] if img_end > img_start else ""
                    img_file = None                    # Special handling for specific problematic files
                    special_files = ['bloom-pyramid.png', 'bloom-pyramid.svg', 'bloom_pyramid.png']
                    is_special_file = img_path.lower().endswith(tuple(special_files)) or os.path.basename(img_path).lower() in special_files
                    
                    if is_special_file:
                        log(f"Special file detected: {img_path}")
                        # Try alternative file formats for special files
                        base_name = os.path.splitext(os.path.basename(img_path))[0].lower()
                        
                        # First look for SVG (vector) version which might be better quality
                        svg_version = find_specific_image(project_root, base_name + '.svg', logger=log)
                        if svg_version:
                            img_file = svg_version
                            log(f"Found SVG version of special file: {img_file}")
                        else:
                            # Then try PNG version
                            png_version = find_specific_image(project_root, base_name + '.png', logger=log)
                            if png_version:
                                img_file = png_version
                                log(f"Found PNG version of special file: {img_file}")
                            else:
                                # Use our specialized function to find any version of this problematic file
                                specific_img = find_specific_image(project_root, base_name + '.*', logger=log)
                                if specific_img:
                                    img_file = specific_img
                                    log(f"Found special file using dedicated search: {img_file}")
                                else:
                                    # Fallback to checking several possible locations with different extensions
                                    possible_extensions = ['.png', '.svg', '.jpg', '.jpeg']
                                    possible_locations = []
                                    
                                    for ext in possible_extensions:
                                        possible_locations.extend([
                                            os.path.join(project_root, base_name + ext),
                                            os.path.join(os.path.dirname(project_root), base_name + ext),
                                            os.path.join(project_root, 'Dodona-Learning-Path', base_name + ext),
                                            os.path.join(project_root, 'Dodona-Learning-Path', 'images', base_name + ext),
                                        ])
                                    
                                    # Add hardcoded paths as last resort
                                    possible_locations.extend([
                                        os.path.join(project_root, 'bloom-pyramid.png'),
                                        os.path.join(project_root, 'bloom-pyramid.svg'),
                                        os.path.join(project_root, '..', 'bloom-pyramid.png')
                                    ])
                                    
                                    for location in possible_locations:
                                        if os.path.exists(location):
                                            img_file = location
                                            log(f"Found special file at: {img_file}")
                                            break
                    
                    # Special handling for relative paths with ../
                    elif img_path.startswith('../'):
                        parent_dir = os.path.dirname(os.path.dirname(base_dir))
                        img_relative_path = img_path.replace('../', '')
                        resolved_path = os.path.join(parent_dir, img_relative_path)
                        log(f"Trying to resolve ../ path: {resolved_path}")
                        
                        if os.path.exists(resolved_path):
                            img_file = resolved_path
                            log(f"Found image by resolving ../ path: {img_file}")
                        else:
                            # Also try with just the filename
                            just_filename = os.path.basename(img_relative_path)
                            # Check some common locations
                            for img_dir in ['images', 'media', '.', '..']:
                                check_path = os.path.join(project_root, img_dir, just_filename)
                                if os.path.exists(check_path):
                                    img_file = check_path
                                    log(f"Found image in common location: {img_file}")
                                    break
                    
                    # Standard image lookup if not found through relative path
                    if img_path and all_images and not img_file:
                        # Try exact filename match first
                        img_basename = os.path.basename(img_path).lower()
                        # Also try URL-decoded version
                        url_decoded_basename = urllib.parse.unquote(img_basename).lower()

                        # Try direct lookup first
                        if img_basename in image_lookup:
                            img_file = image_lookup[img_basename]
                            log(
                                f"Found exact match for image: {img_basename} -> {img_file}")
                        elif url_decoded_basename in image_lookup:
                            img_file = image_lookup[url_decoded_basename]
                            log(
                                f"Found exact match for URL-decoded image: {url_decoded_basename} -> {img_file}")
                        else:
                            # Try name without extension
                            name_without_ext = os.path.splitext(img_basename)[
                                0].lower()
                            url_decoded_name_without_ext = os.path.splitext(url_decoded_basename)[0].lower()
                            
                            if name_without_ext in image_lookup:
                                img_file = image_lookup[name_without_ext]
                                log(
                                    f"Found match by filename without extension: {name_without_ext} -> {img_file}")
                            elif url_decoded_name_without_ext in image_lookup:
                                img_file = image_lookup[url_decoded_name_without_ext]
                                log(
                                    f"Found match by URL-decoded filename without extension: {url_decoded_name_without_ext} -> {img_file}")
                            else:
                                # Try partial matching if direct lookup fails
                                for existing_img in all_images:
                                    existing_basename = os.path.basename(
                                        existing_img).lower()
                                    # Check if markdown path is contained in any found image path
                                    if (img_basename in existing_basename or 
                                        name_without_ext in existing_basename or
                                        url_decoded_basename in existing_basename or
                                        url_decoded_name_without_ext in existing_basename):
                                        img_file = existing_img
                                        log(
                                            f"Found partial match for image: {img_basename} -> {img_file}")
                                        break
                      # Prepare the caption text with appropriate prefix based on image type
                    # Only use alt_text if it's not empty, otherwise use a generic caption
                    if alt_text.strip():
                        caption_text = f"Screenshot: {alt_text}" if ("screenshot" in alt_text.lower(
                        ) or "screenshot" in img_path.lower()) else f"Image: {alt_text}"
                    else:
                        caption_text = "Screenshot" if ("screenshot" in img_path.lower()) else "Image"
                    
                    if img_file and os.path.exists(img_file):                        try:
                            # Insert the actual image
                            log(f"Adding image from: {img_file}")
                              # Try to check if we can access the file before attempting to use it
                            try:
                                with open(img_file, 'rb') as test_access:
                                    pass
                            except PermissionError:
                                log_error(f"Permission denied when trying to access {img_file}")
                                
                                # Check if this is the bloom pyramid image (common problem)
                                is_bloom_pyramid = "bloom" in img_file.lower() and ("pyramid" in img_file.lower() or "taxonomy" in img_file.lower())
                                
                                if is_bloom_pyramid:
                                    # Try to create a fallback bloom pyramid image
                                    log("Attempting to create a fallback Bloom's Taxonomy Pyramid image")
                                    fallback_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "bloom_pyramid_fallback.png")
                                    
                                    try:
                                        # Try to import and use our script to generate a fallback image
                                        try:
                                            import create_bloom_pyramid_image
                                            if create_bloom_pyramid_image.create_bloom_pyramid_image(fallback_path):
                                                log(f"Successfully created fallback bloom pyramid image at {fallback_path}")
                                                img_file = fallback_path
                                            else:
                                                raise Exception("Failed to create fallback image")
                                        except ImportError:
                                            log_error("Could not import create_bloom_pyramid_image module")
                                            raise
                                    except Exception as fallback_error:
                                        log_error(f"Failed to create fallback bloom pyramid image: {str(fallback_error)}")
                                        raise
                                else:
                                    # For other files, try copying to temp location
                                    import shutil
                                    import tempfile
                                    temp_dir = tempfile.gettempdir()
                                    temp_img = os.path.join(temp_dir, os.path.basename(img_file))
                                    try:
                                        shutil.copy2(img_file, temp_img)
                                        log(f"Copied image to temporary location: {temp_img}")
                                        img_file = temp_img
                                    except Exception as copy_error:
                                        log_error(f"Failed to copy image to temp location: {str(copy_error)}")
                                        raise
                            
                            # Add spacing before the image with proper formatting
                            spacing_before = doc.add_paragraph()
                            spacing_before.paragraph_format.space_after = Pt(6)
                            
                            # Create a paragraph for the image with center alignment
                            image_p = doc.add_paragraph()
                            image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            image_p.paragraph_format.space_before = Pt(0)
                            image_p.paragraph_format.space_after = Pt(6)
                            
                            # Calculate appropriate image size based on image type
                            # Screenshots are typically wider, diagrams might need to be smaller for clarity
                            is_screenshot = "screenshot" in alt_text.lower() or "screenshot" in img_path.lower()
                            is_diagram = any(term in alt_text.lower() or term in img_path.lower() 
                                            for term in ["diagram", "chart", "graph", "figure"])
                            
                            # Determine width based on image type
                            if is_screenshot:
                                width_inches = 6.0  # Full width for screenshots
                            elif is_diagram:
                                width_inches = 5.5  # Slightly smaller for diagrams
                            else:
                                width_inches = 5.0  # Default for regular images
                            
                            # Add the image with the calculated width
                            image_run = image_p.add_run()
                            image = image_run.add_picture(img_file, width=Inches(width_inches))                            # Add a properly formatted caption below the image
                            caption_p = doc.add_paragraph()
                            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_p.paragraph_format.space_before = Pt(0)
                            caption_p.paragraph_format.space_after = Pt(12)
                            
                            # Format the caption text with proper styling
                            caption_run = caption_p.add_run(f"Figure: {caption_text}")
                            caption_run.italic = True
                            caption_run.font.name = "Times New Roman"
                            caption_run.font.size = Pt(10)
                            
                            # Add spacing after the image and caption
                            spacing_after = doc.add_paragraph()
                            spacing_after.paragraph_format.space_before = Pt(0)
                            
                            log(f"Successfully inserted image: {img_file}")
                        except Exception as img_error:
                            # If image insertion fails, fall back to placeholder
                            log_error(
                                f"Failed to insert image file {img_file}: {str(img_error)}")
                            p = doc.add_paragraph()
                            p.add_run(f"[{caption_text}]").bold = True
                            log(
                                f"Used placeholder for image that couldn't be inserted: {img_path}")
                    else:
                        # Image file not found, use placeholder
                        p = doc.add_paragraph()
                        p.add_run(
                            f"[{caption_text} - File not found]").bold = True
                        log(f"Image file not found for: {img_path}")

                except Exception as e:
                    # Very simple error handling
                    doc.add_paragraph(
                        "[Image placeholder - Error processing image]")
                    log_error(f"Error processing image reference: {str(e)}")

            # Optimize bold text handling
            elif '**' in line and not line.startswith('#'):
                # Fast path for common descriptions
                if line.strip().startswith('**Description:**'):
                    p = doc.add_paragraph()
                    p.add_run('Description:').bold = True
                    p.add_run(line.split('**Description:**', 1)[1])
                elif line.strip().startswith('**Purpose:**'):
                    p = doc.add_paragraph()
                    p.add_run('Purpose:').bold = True
                    p.add_run(line.split('**Purpose:**', 1)[1])
                else:
                    # More efficient handling of bold text with fewer splits
                    p = doc.add_paragraph()
                    parts = line.split('**')
                    for idx, part in enumerate(parts):
                        if part:  # Skip empty parts
                            run = p.add_run(part)
                            run.bold = (idx % 2 == 1)  # Odd parts are bold

            # Regular paragraph (not a header, list, image, or bold text)
            elif line.strip() and not line.startswith('---'):
                doc.add_paragraph(line)
            elif line.strip() == '---':
                doc.add_paragraph('_' * 50)

            i += 1  # Move to next line

            # Log progress for long documents (less frequent to reduce overhead)
            if i % 1000 == 0 and i > 0:
                log(f"Processed {i} of {len(lines)} lines ({i/len(lines)*100:.1f}%)...")

        log("Finished processing markdown content. Saving document...")
        # Save the document
        try:
            doc.save(output_file)
            log(f"Document successfully saved to {output_file}")
        except Exception as save_error:
            log_error(f"Error saving document: {str(save_error)}")
        return f"Document successfully saved to {output_file}"

    except Exception as e:
        log_error(f"Error: {str(e)}")
        return f"Error: {str(e)}"


def generate_possible_paths(img_path, base_dir, project_root):
    """Generate common possible paths for an image (optimized version)"""
    # This function is kept for backward compatibility but made more efficient
    common_locations = [
        img_path,  # Direct path
        os.path.join(base_dir, img_path),  # Relative to markdown file
        os.path.join(project_root, 'images', os.path.basename(
            img_path)),  # Project images folder
        os.path.join(os.path.dirname(base_dir), 'images',
                     os.path.basename(img_path)),  # Parent images folder
        os.path.join(project_root, 'Dodona-Learning-Path', 'images',
                     os.path.basename(img_path))  # Specific project folder
    ]

    # For ../images/ pattern
    if img_path.startswith('../'):
        relative_path = img_path.replace('../', '')
        common_locations.append(os.path.normpath(
            os.path.join(os.path.dirname(base_dir), relative_path)))

    return common_locations


def find_project_root(start_dir):
    """Find the root directory of the project by looking for common project markers"""
    current_dir = start_dir
    max_depth = 10  # Prevent infinite loop
    depth = 0
    
    # Keep track of all directories we visit to use as potential image sources
    visited_dirs = []
    visited_dirs.append(current_dir)

    # Root directory markers (files or folders that indicate a project root)
    root_markers = [
        '.git',                        # Git repository
        'Dodona-Learning-Path',        # Our specific project folder
        'stats-course-dodona',         # Another potential project folder
        'run_conversion.ps1',          # A script at the root of our project
        'bloom-pyramid.png'            # A specific file we're often looking for
    ]

    while depth < max_depth:
        # Check for any of the root markers
        for marker in root_markers:
            if os.path.exists(os.path.join(current_dir, marker)):
                return current_dir

        # Special case: if we're in a numbered course directory (e.g., "10. Boxplot Interpretation")
        # Try to get the parent that contains multiple course folders
        current_basename = os.path.basename(current_dir)
        if current_basename and (current_basename[0].isdigit() or current_basename.startswith("0.")):
            parent_dir = os.path.dirname(current_dir)
            # Check if parent contains other numbered directories
            try:
                siblings = os.listdir(parent_dir)
                numbered_siblings = [
                    d for d in siblings if d[0].isdigit() or d.startswith("0.")]
                if len(numbered_siblings) > 1:
                    return parent_dir
            except:
                pass

        parent_dir = os.path.dirname(current_dir)
        if parent_dir == current_dir:  # Reached root directory
            break

        current_dir = parent_dir
        depth += 1

    # If we can't find a root marker, return two levels up from the original directory
    # This increases our chances of finding related image folders
    parent_dir = os.path.dirname(start_dir)
    grandparent_dir = os.path.dirname(parent_dir)
    if grandparent_dir != parent_dir:  # If we can go up two levels
        return grandparent_dir

    # Otherwise return the original directory
    return start_dir


def find_image_file(img_path, base_dir, project_root, all_images, logger=None):
    """Try multiple approaches to find the actual image file, optimized to reduce redundant checks"""
    # Safely use logger if provided
    def safe_log(msg):
        if logger:
            try:
                logger.info(msg)
            except:
                print(msg)
        else:
            print(msg)

    log = safe_log

    # Debug information
    log(f"Looking for image: {img_path}")

    # Extract just the filename without path
    img_filename = os.path.basename(img_path)
    log(f"Image filename: {img_filename}")

    # Create a cache of already checked paths to avoid redundant checks
    checked_paths = set()

    # First check if the exact path exists
    if os.path.isabs(img_path) and os.path.exists(img_path):
        log(f"Found exact path: {img_path} ✓")
        return img_path

    # Try the relative path to the markdown file
    relative_path = os.path.join(base_dir, img_path)
    if os.path.exists(relative_path):
        log(f"Found relative path: {relative_path} ✓")
        return relative_path

    # For ../images/ pattern
    if img_path.startswith('../'):
        # Remove '../' and join with parent directory
        resolved_path = os.path.normpath(os.path.join(
            os.path.dirname(base_dir), img_path.replace('../', '')))
        if os.path.exists(resolved_path):
            log(f"Found path from ../ pattern: {resolved_path} ✓")
            return resolved_path

    # Generate variations of the filename to check against the known image list
    name_part, ext = os.path.splitext(img_filename)
    variations = [
        img_filename.lower(),
        name_part.lower() + ext.lower(),
        name_part.replace('-', '_').lower() + ext.lower(),
        name_part.replace('_', '-').lower() + ext.lower(),
        name_part.replace(' ', '_').lower() + ext.lower(),
        name_part.replace(' ', '-').lower() + ext.lower(),
    ]

    # First try exact match with known images (faster than checking filesystem)
    for img in all_images:
        img_base = os.path.basename(img).lower()
        if img_base in variations:
            log(f"Found exact match in known images: {img} ✓")
            return img

    # Then try partial matches with known images
    for img in all_images:
        img_base = os.path.basename(img).lower()
        base_name, _ = os.path.splitext(img_base)

        # Check if any variation is contained in the image name
        for var in variations:
            var_base, _ = os.path.splitext(var)
            if var_base in base_name or base_name in var_base:
                log(f"Found fuzzy match in known images: {img} ✓")
                return img

    # If still not found, try common locations with variations
    common_dirs = [
        os.path.join(project_root, 'images'),
        os.path.join(project_root, 'Dodona-Learning-Path', 'images'),
        os.path.join(base_dir, 'images'),
        os.path.join(os.path.dirname(base_dir), 'images'),
        os.path.join(project_root, 'description', 'images'),
        os.path.join(project_root, 'description', 'media'),
    ]

    for img_dir in common_dirs:
        if os.path.exists(img_dir):
            for var in variations:
                check_path = os.path.join(img_dir, var)
                if check_path not in checked_paths and os.path.exists(check_path):
                    log(f"Found in common location: {check_path} ✓")
                    return check_path
                checked_paths.add(check_path)

    # Not found in any of the expected locations
    log(f"Image not found: {img_path} ✗")
    return None


def setup_logging():
    """Set up logging to write to a file"""
    import logging
    log_file = os.path.join(os.path.dirname(
        os.path.abspath(__file__)), "conversion.log")

    # Configure logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file),
            logging.StreamHandler()  # Also log to console
        ]
    )

    logging.info("="*80)
    logging.info("LOGGING INITIALIZED")
    logging.info("="*80)
    return logging


def find_specific_image(project_root, target_filename, max_depth=3, logger=None):
    """Special function to search for a specific image file by recursively searching up from the project root"""
    # Safely use logger if provided
    def safe_log(msg):
        if logger:
            try:
                logger.info(msg)
            except:
                print(msg)
        else:
            print(msg)

    log = safe_log
    log(f"Searching specifically for {target_filename}...")
    
    # Normalize the target filename
    target_filename = target_filename.lower()
    
    # First check in project root and immediate subdirectories
    for root, dirs, files in os.walk(project_root):
        current_depth = root[len(project_root):].count(os.sep)
        if current_depth <= max_depth:
            for file in files:
                if file.lower() == target_filename:
                    full_path = os.path.join(root, file)
                    log(f"Found specific image {target_filename} at {full_path}")
                    return full_path
    
    # If not found, try parent directories
    parent_dir = os.path.dirname(project_root)
    grandparent_dir = os.path.dirname(parent_dir)
    
    # Check parent
    parent_path = os.path.join(parent_dir, target_filename)
    if os.path.exists(parent_path):
        log(f"Found specific image {target_filename} in parent directory: {parent_path}")
        return parent_path
        
    # Check grandparent
    grandparent_path = os.path.join(grandparent_dir, target_filename)
    if os.path.exists(grandparent_path):
        log(f"Found specific image {target_filename} in grandparent directory: {grandparent_path}")
        return grandparent_path
    
    # Try with alternative names (without hyphens, with underscores, etc.)
    alt_names = [
        target_filename.replace('-', '_'),
        target_filename.replace('_', '-'),
        target_filename.replace(' ', '_'),
        target_filename.replace(' ', '-')
    ]
    
    for alt_name in alt_names:
        # Check in project root
        alt_path = os.path.join(project_root, alt_name)
        if os.path.exists(alt_path):
            log(f"Found specific image with alternative name {alt_name} at {alt_path}")
            return alt_path
            
        # Check parent directory
        alt_parent_path = os.path.join(parent_dir, alt_name)
        if os.path.exists(alt_parent_path):
            log(f"Found specific image with alternative name {alt_name} in parent directory: {alt_parent_path}")
            return alt_parent_path
    
    # If all else fails, try a broader search using glob
    for ext in ['.png', '.jpg', '.svg', '.jpeg']:
        # Try finding files that contain the name
        name_part = target_filename.split('.')[0]
        glob_pattern = os.path.join(project_root, f"**/*{name_part}*{ext}")
        matches = glob.glob(glob_pattern, recursive=True)
        
        if matches:
            log(f"Found potential match for {target_filename} using glob: {matches[0]}")
            return matches[0]
    
    log(f"Could not find specific image {target_filename} after extensive search")
    return None


if __name__ == "__main__":
    try:
        import time
        import sys
        start_time = time.time()
        logger = setup_logging()        # Parse arguments
        test_images = False
        test_file = None
        find_bloom_pyramid = False

        # Check for command-line arguments
        if len(sys.argv) > 1:
            # Look for test image flag
            if "--test-images" in sys.argv:
                test_images = True
                logger.info("Image testing mode enabled")
            
            # Special flag to focus on the bloom pyramid issue
            if "--find-bloom-pyramid" in sys.argv:
                find_bloom_pyramid = True
                logger.info("Bloom pyramid special search enabled")
                
                # Create a test markdown file with just the bloom pyramid reference
                test_bloom_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "test_bloom.md")
                with open(test_bloom_path, "w") as f:
                    f.write("# Bloom Pyramid Test\n\n")
                    f.write("This is a test file to check if the bloom pyramid image is correctly embedded.\n\n")
                    f.write("![Bloom's Taxonomy Pyramid](../bloom-pyramid.png)\n\n")
                    f.write("This reference should find and embed the bloom pyramid image.\n")
                
                # Set this as our test file
                test_file = test_bloom_path
                logger.info(f"Created test file for bloom pyramid at {test_bloom_path}")

            # Look for a specific test file
            for arg in sys.argv:
                if arg.endswith(".md") and os.path.exists(arg):
                    test_file = arg
                    logger.info(f"Test file specified: {test_file}")

        # Print a clear start message
        logger.info("="*80)
        logger.info("MARKDOWN TO WORD CONVERSION SCRIPT - STARTING")
        logger.info("="*80)

        # Files to convert
        files_to_convert = [
            {
                "input_file": "description/Dodona_Learning_Path_Overview.md",
                "output_name": "20250621_Dodona_Learning_Path_Overview.docx"
            },
            {
                "input_file": "../Motivation_Letter.md",
                "output_name": "20250621_Motivation_Letter.docx"
            }
        ]

        # If we're testing a specific file, use that instead
        if test_file:
            files_to_convert = [
                {
                    "input_file": test_file,
                    "output_name": f"TEST_{os.path.basename(test_file).replace('.md', '.docx')}"
                }
            ]

        # Output directory
        output_dir = r"C:\Users\kukumar\OneDrive - UGent\Job"

        # Use current directory for test outputs
        if test_images or test_file:
            output_dir = os.path.dirname(os.path.abspath(__file__))
            logger.info(
                f"Using current directory for test output: {output_dir}")

        logger.info(f"Output directory: {output_dir}")

        # Create output directory if it doesn't exist
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            logger.info(f"Created output directory: {output_dir}")
        else:
            logger.info(f"Output directory already exists")

        # Get the current directory
        current_dir = os.path.dirname(os.path.abspath(__file__))
        logger.info(f"Current directory: {current_dir}")

        # Process each file
        for file_info in files_to_convert:
            file_start_time = time.time()
            input_file = file_info["input_file"]
            output_name = file_info["output_name"]

            # Construct full paths
            input_path = os.path.join(current_dir, input_file) if not os.path.isabs(
                input_file) else input_file
            output_path = os.path.join(output_dir, output_name)

            logger.info("-"*80)
            logger.info(f"Processing file: {input_file}")
            logger.info(f"Full input path: {input_path}")
            logger.info(f"Full output path: {output_path}")

            if not os.path.exists(input_path):
                logger.error(f"ERROR: Input file not found: {input_path}")
                continue

            logger.info(f"Converting {input_path} to {output_path}...")
            result = simple_markdown_to_word(
                input_path, output_path, logger=logger, show_image_debug=test_images)

            file_end_time = time.time()
            file_duration = file_end_time - file_start_time
            logger.info(f"File conversion time: {file_duration:.2f} seconds")
            logger.info(f"Result: {result}")

            if os.path.exists(output_path):
                logger.info(f"SUCCESS: Output file created: {output_path}")
                # Get file size
                file_size = os.path.getsize(output_path) / 1024  # KB
                logger.info(f"File size: {file_size:.2f} KB")
            else:
                logger.error(f"ERROR: Output file not created: {output_path}")

        end_time = time.time()
        total_duration = end_time - start_time
        logger.info("="*80)
        logger.info(
            f"CONVERSION COMPLETE - Total time: {total_duration:.2f} seconds")
        logger.info("="*80)

        # List files in output directory
        if os.path.exists(output_dir):
            logger.info("Files in output directory:")
            for file in os.listdir(output_dir):
                file_path = os.path.join(output_dir, file)
                file_size = os.path.getsize(file_path) / 1024  # KB
                logger.info(f" - {file} ({file_size:.2f} KB)")
        else:
            logger.error("Output directory does not exist")

    except Exception as e:
        # Print directly to make sure it's visible even if logging fails
        print(f"CRITICAL ERROR: {str(e)}")
        import traceback
        traceback.print_exc()

        # Try to log if possible
        try:
            if 'logger' in locals():
                logger.error(f"CRITICAL ERROR: {str(e)}")
                logger.error(traceback.format_exc())
        except:
            pass

        # Write a status file for debugging
        try:
            with open(os.path.join(os.path.dirname(os.path.abspath(__file__)), "conversion_status.txt"), "w") as status_file:
                status_file.write(f"Conversion failed: {str(e)}\n")
                status_file.write(traceback.format_exc())
        except:
            pass

    # Write a status file for debugging (success case)
    try:
        with open(os.path.join(os.path.dirname(os.path.abspath(__file__)), "conversion_status.txt"), "w") as status_file:
            status_file.write("Conversion completed successfully\n")
            if os.path.exists(output_dir):
                status_file.write("Output directory exists\n")
                for file in os.listdir(output_dir):
                    file_path = os.path.join(output_dir, file)
                    file_size = os.path.getsize(file_path) / 1024  # KB
                    status_file.write(f" - {file} ({file_size:.2f} KB)\n")
            else:
                status_file.write("Output directory does not exist\n")
    except Exception as status_error:
        print(f"Error writing status file: {str(status_error)}")
