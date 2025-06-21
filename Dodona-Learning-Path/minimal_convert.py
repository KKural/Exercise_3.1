"""
Simple markdown to Word converter - Minimal version with no image handling
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

# Set up logging
log_file = os.path.join(os.path.dirname(
    os.path.abspath(__file__)), "minimal_conversion.log")
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file),
        logging.StreamHandler()  # Also log to console
    ]
)
logger = logging.getLogger()


def minimal_markdown_to_word(md_file, output_file):
    """Convert markdown to Word without image handling"""
    try:
        logger.info(f"Starting conversion of {md_file}")

        # Read the markdown file
        with open(md_file, 'r', encoding='utf-8') as file:
            md_content = file.read()

        # Create a new Word document
        doc = Document()

        # Set font for the entire document
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)

        # Process content as simple paragraphs
        lines = md_content.split('\n')

        for line in lines:
            # Skip empty lines
            if not line.strip():
                continue

            # Simple header processing
            if line.startswith('# '):
                doc.add_heading(line[2:], 0)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 2)
            # Skip image lines
            elif line.strip().startswith('!['):
                doc.add_paragraph('[Image placeholder]')
            # Regular paragraph
            else:
                doc.add_paragraph(line)

        # Save the document
        doc.save(output_file)
        logger.info(f"Successfully saved document to {output_file}")
        return True

    except Exception as e:
        logger.error(f"Error in conversion: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return False


if __name__ == "__main__":
    logger.info("="*80)
    logger.info("MINIMAL MARKDOWN TO WORD CONVERTER - STARTING")
    logger.info("="*80)

    # Files to convert
    files_to_convert = [
        {
            "input_file": "description/Dodona_Learning_Path_Overview.md",
            "output_name": "20250621_Dodona_Learning_Path_Overview.docx"
        },
        {
            "input_file": "../Motivation_Letter.md",
            "output_name": "20250621_Motivation_Letter.docx"
        }
    ]    # Output directory - use a local directory instead of network path
    output_dir = os.path.join(os.path.dirname(
        os.path.abspath(__file__)), "output")
    logger.info(f"Output directory: {output_dir}")

    # Create output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        logger.info(f"Created output directory: {output_dir}")
    else:
        logger.info(f"Output directory already exists")

    # Get the current directory
    current_dir = os.path.dirname(os.path.abspath(__file__))
    logger.info(f"Current directory: {current_dir}")

    # Process each file
    for file_info in files_to_convert:
        input_file = file_info["input_file"]
        output_name = file_info["output_name"]

        # Construct full paths
        input_path = os.path.join(current_dir, input_file)
        output_path = os.path.join(output_dir, output_name)

        logger.info("-"*80)
        logger.info(f"Processing file: {input_file}")
        logger.info(f"Full input path: {input_path}")
        logger.info(f"Full output path: {output_path}")

        if not os.path.exists(input_path):
            logger.error(f"ERROR: Input file not found: {input_path}")
            continue

        logger.info(f"Converting {input_path} to {output_path}...")
        success = minimal_markdown_to_word(input_path, output_path)

        if success and os.path.exists(output_path):
            logger.info(f"SUCCESS: Output file created: {output_path}")
            # Get file size
            file_size = os.path.getsize(output_path) / 1024  # KB
            logger.info(f"File size: {file_size:.2f} KB")
        else:
            logger.error(f"ERROR: Output file not created: {output_path}")

    logger.info("="*80)
    logger.info("CONVERSION COMPLETE")
    logger.info("="*80)

    # Write a status file
    with open(os.path.join(os.path.dirname(os.path.abspath(__file__)), "minimal_conversion_status.txt"), "w") as status_file:
        status_file.write("Minimal conversion completed\n")
        if os.path.exists(output_dir):
            status_file.write("Output directory exists\n")
            for file in os.listdir(output_dir):
                file_path = os.path.join(output_dir, file)
                file_size = os.path.getsize(file_path) / 1024  # KB
                status_file.write(f" - {file} ({file_size:.2f} KB)\n")
        else:
            status_file.write("Output directory does not exist\n")
