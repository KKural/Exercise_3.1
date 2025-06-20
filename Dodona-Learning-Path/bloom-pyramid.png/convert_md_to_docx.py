from bs4 import BeautifulSoup
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor  # Added RGBColor
from docx import Document
import markdown
import os
import sys

# Check for required packages and install them if they're missing
required_packages = ['python-docx', 'markdown', 'beautifulsoup4']
missing_packages = []

for package in required_packages:
    try:
        __import__(package.replace('-', '_'))
    except ImportError:
        missing_packages.append(package)

if missing_packages:
    print(f"Missing required packages: {', '.join(missing_packages)}")
    print("Please run install_requirements.py first or install them manually with:")
    print(f"pip install {' '.join(missing_packages)}")
    sys.exit(1)


def markdown_to_word(md_file, output_file):
    try:
        # Read the markdown file
        with open(md_file, 'r', encoding='utf-8') as file:
            md_content = file.read()

        print(f"Successfully read Markdown file: {md_file}")

        # Create a new Word document
        doc = Document()

        # Set font for the entire document
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'  # Changed from Calibri to Times New Roman
        style.font.size = Pt(11)
        
        # Update heading styles to use black color
        for i in range(1, 10):  # Update all heading levels
            if f'Heading {i}' in doc.styles:
                heading_style = doc.styles[f'Heading {i}']
                heading_style.font.name = 'Times New Roman'
                heading_style.font.color.rgb = RGBColor(0, 0, 0)  # Black color

        # Convert markdown to HTML
        html = markdown.markdown(md_content, extensions=[
                                 'tables', 'fenced_code'])

        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')

        # Get the base directory for image resolution
        base_dir = os.path.dirname(os.path.abspath(md_file))

        # Process the document
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'table', 'img']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                doc.add_heading(element.text, level)

            elif element.name == 'p':
                # Check if this paragraph contains an image
                img = element.find('img')
                if img:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_src = img.get('src', '')
                    alt_text = img.get('alt', 'Image')

                    # Resolve relative path
                    if img_src.startswith('../'):
                        # Remove '../' and join with parent directory
                        relative_path = img_src.replace('../', '')
                        img_full_path = os.path.normpath(os.path.join(
                            os.path.dirname(base_dir), relative_path))
                    else:
                        img_full_path = os.path.normpath(
                            os.path.join(base_dir, img_src))

                    # Check if image exists
                    if os.path.exists(img_full_path):
                        try:
                            doc.add_picture(img_full_path, width=Inches(6))
                            doc.add_paragraph(f"{alt_text}").italic = True
                        except Exception as e:
                            doc.add_paragraph(
                                f"[Image: {alt_text} - {img_src}]\nError: {str(e)}").italic = True
                    else:
                        doc.add_paragraph(
                            f"[Image: {alt_text} - {img_src}]\nFile not found: {img_full_path}").italic = True
                else:
                    doc.add_paragraph(element.text)

            elif element.name == 'table':
                # Get number of rows and columns
                rows = element.find_all('tr'):
                if not rows:, style='List Bullet')
                    continue
e == 'ol':
                # Get the maximum number of cells in any row                for li in element.find_all('li'):
                max_cells = max(umber')
                    len(row.find_all(['td', 'th'])) for row in rows)

                # Create the table                # Get number of rows and columns
                table = doc.add_table(rows=len(rows), cols=max_cells)d_all('tr')
                table.style = 'Table Grid'

                # Fill the table
                for i, row in enumerate(rows):m number of cells in any row
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):r row in rows)
                        table.cell(i, j).text = cell.text.strip()

        # Save the document                table = doc.add_table(rows=len(rows), cols=max_cells)
        doc.save(output_file) = 'Table Grid'
        return f"Document successfully saved to {output_file}"

    except Exception as e:                for i, row in enumerate(rows):
        return f"Error: {str(e)}"= row.find_all(['td', 'th'])
n enumerate(cells):
                        table.cell(i, j).text = cell.text.strip()
# Usage
if __name__ == "__main__": # Save the document
    # This script is in the images directory, so we need to adjust the path for the input filele)
    input_file = "../Dodona_Learning_Path_Overview.md"
    output_file = "../Dodona_Learning_Path_Overview.docx"

    # Get the current directory        return f"Error: {str(e)}"
    current_dir = os.path.dirname(os.path.abspath(__file__))

    # Construct full paths# Usage
    input_path = os.path.normpath(os.path.join(current_dir, input_file))
    output_path = os.path.normpath(os.path.join(current_dir, output_file))ath for the input file

    if not os.path.exists(input_path):    output_file = "../Dodona_Learning_Path_Overview.docx"
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)
s.path.dirname(os.path.abspath(__file__))
    print(f"Converting {input_path} to {output_path}...")
    result = markdown_to_word(input_path, output_path)
    print(result)_dir, input_file))
    output_path = os.path.normpath(os.path.join(current_dir, output_file))

    if not os.path.exists(input_path):
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)

    print(f"Converting {input_path} to {output_path}...")
    result = markdown_to_word(input_path, output_path)
    print(result)
