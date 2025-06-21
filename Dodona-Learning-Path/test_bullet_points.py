import os
import re
import time
import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')


def convert_md_to_docx(md_path, output_path, debug_mode=False):
    try:
        if debug_mode:
            logging.info(f"Debug mode: Processing {md_path} to {output_path}")

        # Create a test file with the special bullet point format
        test_md_content = """# Test Document

## Test Section

### Test Subsection

- **Correct (A)**: This is a correct answer with an explanation.
- **Incorrect (B)**: This is an incorrect answer with explanation.
"""

        # Write the test content to a temporary file
        test_md_path = os.path.join(os.path.dirname(md_path), "temp_test.md")
        with open(test_md_path, 'w', encoding='utf-8') as f:
            f.write(test_md_content)

        # Process the test file
        start_time = time.time()
        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Set font for the entire document
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)

        # Set paragraph spacing
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.15

        # Update heading styles for better formatting
        for i in range(1, 10):
            if f'Heading {i}' in doc.styles:
                heading_style = doc.styles[f'Heading {i}']
                heading_style.font.name = 'Times New Roman'
                heading_style.font.color.rgb = RGBColor(0, 0, 0)
                heading_style.font.bold = True

                # Add proper spacing for headings
                heading_style.paragraph_format.space_before = Pt(12)
                heading_style.paragraph_format.space_after = Pt(6)

                # Set specific size based on heading level
                if i == 1:
                    heading_style.font.size = Pt(16)
                elif i == 2:
                    heading_style.font.size = Pt(14)
                else:
                    heading_style.font.size = Pt(13)

        # Read the content of the test markdown file
        with open(test_md_path, 'r', encoding='utf-8') as f:
            md_content = f.read().split('\n')

        # Process each line of the markdown
        current_list = []
        in_list = False
        in_table = False

        for line in md_content:
            # Handle headings (h1 to h6)
            if line.startswith('#'):
                level = len(line.split(' ')[0])  # Count the number of #
                heading_text = line.strip('#').strip()
                doc.add_heading(heading_text, level=level)

            # Check for bullet points and numbered lists
            elif line.strip().startswith(('-', '*', '1.')):
                # Extract the content without the bullet marker
                content = line.strip()[2:].strip() if line.strip().startswith(
                    ('-', '*')) else line.strip()[3:].strip()

                # Check if the bullet point contains bold text
                if '**' in content:
                    p = doc.add_paragraph('', style='List Bullet' if line.strip(
                    ).startswith(('-', '*')) else 'List Number')

                    # Process bold formatting within the bullet point
                    parts = content.split('**')
                    for idx, part in enumerate(parts):
                        if part:  # Skip empty parts
                            run = p.add_run(part)
                            run.bold = (idx % 2 == 1)  # Odd parts are bold
                else:
                    # Standard bullet point without special formatting
                    doc.add_paragraph(content, style='List Bullet' if line.strip(
                    ).startswith(('-', '*')) else 'List Number')

            # Handle blank line
            elif not line.strip():
                pass  # Skip blank lines

            # Default paragraph
            else:
                doc.add_paragraph(line)

        # Save the document
        test_output_path = os.path.join(
            os.path.dirname(output_path), "test_output.docx")
        doc.save(test_output_path)

        # Clean up
        os.remove(test_md_path)

        logging.info(
            f"Test completed successfully. Output saved to {test_output_path}")
        return True

    except Exception as e:
        logging.error(f"Error in test: {e}")
        return False


if __name__ == "__main__":
    # Path to the workspace
    workspace_path = r"c:\Users\kukumar\OneDrive - UGent\Desktop\basic-stats-dodona\stats-course-dodona\stats-course-dodona"

    # Test paths
    test_md_path = os.path.join(
        workspace_path, "Dodona-Learning-Path", "description", "test.md")
    test_output_path = os.path.join(
        workspace_path, "Dodona-Learning-Path", "description", "test_output.docx")

    # Run the test
    convert_md_to_docx(test_md_path, test_output_path, debug_mode=True)
